from docx import Document
from docx.shared import Inches
import os

class DocxLibrary:
    """Dynamic DOCX library - FIXED for Robot Framework lists."""
    
    def create_screenshots_document(self, output_path, *args, HeaderReport):
        """Handles Robot Framework @{list} unpacking correctly."""
        print(f"DEBUG: Received {len(args)} raw arguments")
        
        # FLATTEN Robot Framework nested lists
        flattened_args = []
        for arg in args:
            if isinstance(arg, list):
                flattened_args.extend(arg)  # Flatten nested lists
            else:
                flattened_args.append(arg)
        
        args = tuple(flattened_args)
        print(f"DEBUG: Flattened to {len(args)} arguments ({len(args)//2} screenshots)")
        
        if len(args) % 2 != 0:
            raise AssertionError(f"Expected even number of args. Got {len(args)}")
        
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        
        doc = Document()
        doc.add_heading(HeaderReport, 0)
        # doc.add_paragraph(f'Generated: {os.path.basename(output_path)}')
        doc.add_paragraph('')
        
        for i in range(0, len(args), 2):
            img_path = args[i]
            caption = args[i + 1]
            
            print(f"  📸 Adding: {os.path.basename(img_path)} - {caption}")
            
            if not os.path.exists(img_path):
                raise AssertionError(f"❌ Image missing: {img_path}")
            
            doc.add_heading(caption, level=1)
            doc.add_picture(img_path, width=Inches(6.5))
            doc.add_paragraph('')
        
        doc.save(output_path)
        print(f"✅ SUCCESS: DOCX saved to {output_path}")
        return output_path
